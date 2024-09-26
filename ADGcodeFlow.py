import os
import streamlit as st
from langchain_openai import AzureChatOpenAI
from langchain_openai import AzureOpenAIEmbeddings
#from ollama_modell import process_file_with_ollama  # Ensure this import is correct
from langchain_community.llms import Ollama
from langchain.document_loaders import UnstructuredFileLoader
from langchain_community.vectorstores import FAISS
#from langchain.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.chains import RetrievalQA
import docx
from docx import Document
from dotenv import load_dotenv
from io import BytesIO
load_dotenv() 

os.environ["AZURE_OPENAI_API_VERSION"] = st.secrets["AZURE_OPENAI_API_VERSION"]
os.environ["AZURE_OPENAI_CHAT_DEPLOYMENT_NAME"] = st.secrets["AZURE_OPENAI_CHAT_DEPLOYMENT_NAME"]
os.environ["AZURE_OPENAI_API_KEY"] = st.secrets["AZURE_OPENAI_API_KEY"]
os.environ["AZURE_OPENAI_ENDPOINT"] = st.secrets["AZURE_OPENAI_ENDPOINT"]
# os.environ["http_proxy"] = st.secrets["http_proxy"]
# os.environ["https_proxy"] = st.secrets["https_proxy"]

def process_file_with_ollama(file_path, question):
    # Load the LLM
    llm = AzureChatOpenAI(openai_api_version=os.environ["AZURE_OPENAI_API_VERSION"],
                      azure_deployment=os.environ["AZURE_OPENAI_CHAT_DEPLOYMENT_NAME"],
                      temperature=0.1)

    # Load the document dynamically based on the file_path
    loader = UnstructuredFileLoader(file_path)
    documents = loader.load()

    # Split the text into chunks
    text_splitter = CharacterTextSplitter(separator="/n",
                                          chunk_size=1000,
                                          chunk_overlap=200)
    text_chunks = text_splitter.split_documents(documents)

    # Load the vector embedding model
    # embeddings = HuggingFaceEmbeddings(
    # model_name="sentence-transformers/all-MiniLM-L6-v2"
    # )
    embeddings = AzureOpenAIEmbeddings(
        azure_deployment="aldmembeddings",
        openai_api_version="2023-06-01-preview",
    )

    # Create the FAISS knowledge base
    knowledge_base = FAISS.from_documents(text_chunks, embeddings)

    # Create the Retrieval QA chain
    qa_chain = RetrievalQA.from_chain_type(
        llm,
        retriever=knowledge_base.as_retriever()
    )

    # Process the question and get a response
    response = qa_chain.invoke({"query": question})
    
    # Return the response
    return response["result"]

# file_path = 'HLS - CR DEX7629-8 - Fast Data Solution v.2_0 (ENG).docx'
# doc = docx.Document(file_path)


# def create_docx(template_path, text):
#             doc = Document(template_path)
            
#             for paragraph in doc.paragraphs:
#                 if 'TEMPLATE' in paragraph.text:
#                     paragraph.text = paragraph.text.replace('TEMPLATE', text)
                    
#             bio = BytesIO()
#             doc.save(bio)
#             return bio.getvalue()
        

# Streamlit app title
st.title("Code Flow Analyzer using LLM")

# File uploader widget
uploaded_file = st.file_uploader("Upload a Python file", type=["py", "txt"])

if uploaded_file is not None:
    # Ensure the "temp" directory exists
    temp_dir = "temp"
    if not os.path.exists(temp_dir):
        os.makedirs(temp_dir)  # Create the directory if it doesn't exist
    
    # Save the uploaded file to the "temp" directory
    file_path = os.path.join(temp_dir, uploaded_file.name)
    
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    
    st.write(f"File {uploaded_file.name} uploaded successfully!")

    # Default question for analysis
    question = """Analyze the following Python code and generate a structured documentation flow. Include the project/module name, an overview, dependencies, installation instructions, detailed function/class documentation, code examples, license information, authors, changelog, and any additional notes. Follow this format:
Project/Module Name:
[Insert name here]
Overview:
A brief description of the purpose and functionality of the project or module.
Dependencies:
List any external libraries, frameworks, or systems that the code depends on.
Installation:
Instructions on how to set up the project or module.
bash
# Example: pip install <dependency> THIS 

Function/Class Documentation:
Function/Class Name:
A clear, descriptive name of the function or class.
Description: A short summary of what the function or class does.
Parameters:
param1 (type): Description of the first parameter.
param2 (type): Description of the second parameter.
Returns: Type of the return value and what it represents.
Raises/Exceptions: Any errors or exceptions that might be raised.
Example Code:
python
# Example
def add(a: int, b: int) -> int:
    \"\"\"Adds two integers and returns the result.\"\"\"
    return a + b

# Example usage
result = add(3, 5)  # result is 8

License:
Include details about the licensing of the project or module.
Authors:
List the authors or contributors to the project.
Changelog:
A log of changes made to the codebase over time.
Version 1.0.0 - Initial release
Additional Notes:
Any other relevant information or tips for users."""

    # Run the Ollama model to process the file and display the result
if st.button("Generate flow"):
    with st.spinner("Processing..."):
        response = process_file_with_ollama(file_path, question)
        st.subheader("Code Flow Description")
        st.write(response)
        
        # template_path = 'HLS - CR DEX7629-8 - Fast Data Solution v.2_0 (ENG).docx'
        
        # st.download_button(
        #         label="Download result as DOCX",
        #         data=create_docx(template_path, response),
        #         file_name="code_flow_description.docx",
        #         mime="docx",
        #     )    
